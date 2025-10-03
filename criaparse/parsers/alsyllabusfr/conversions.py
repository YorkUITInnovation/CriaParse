import io
import re
from typing import List, Tuple, Any

import mammoth
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from docx.text.paragraph import Paragraph

from criaparse.models import ElementType, Element


def find_hlevel(doc: Document) -> List[str]:
    headings = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            headings.append(paragraph.text.strip())
    sections = [item for item in headings if item]
    return sections


def find_sections_paragraphs(sections: List[str], doc: Document) -> List[int]:
    section_paragraphs = []
    for i in range(len(sections)):
        for j in range(len(doc.paragraphs)):
            if sections[i] == doc.paragraphs[j].text.strip():
                section_paragraphs.append(j)
    return section_paragraphs


def include_hyperlink(paragraph: Paragraph) -> Tuple[List[str], List[str]]:
    hyperlink_text = []
    hyperlink_url = []
    if len(paragraph.hyperlinks) > 0:
        for hyperlink in paragraph.hyperlinks:
            hyperlink_text.append(hyperlink.text)
            hyperlink_url.append(hyperlink.url)
    return hyperlink_text, hyperlink_url


def convert_doc_to_nodes(section_paragraphs: List[int], doc: Document, sections: List[str]) -> List[str]:
    nodes_text = []
    nodes_temp = ""
    for i in range(len(section_paragraphs) - 1):
        for j in range(section_paragraphs[i] + 1, section_paragraphs[i + 1]):
            hyperlink_text, hyperlink_url = include_hyperlink(doc.paragraphs[j])
            for k in range(len(hyperlink_text)):
                if len(hyperlink_text) > 0:
                    temp_text = doc.paragraphs[j].text.replace(hyperlink_text[k],
                                                               "[" + hyperlink_text[k] + "](" + hyperlink_url[k] + ")")
                    doc.paragraphs[j].text = temp_text
            nodes_temp = nodes_temp + doc.paragraphs[j].text.strip() + " "
        nodes_text.append("*" + sections[i] + "*\n" + nodes_temp + "\n")
        nodes_temp = ""
    for i in range(section_paragraphs[len(section_paragraphs) - 1], len(doc.paragraphs)):
        if i == section_paragraphs[len(section_paragraphs) - 1]:
            nodes_temp = nodes_temp + "*" + doc.paragraphs[i].text.strip() + "*\n"
        else:
            hyperlink_text, hyperlink_url = include_hyperlink(doc.paragraphs[i])
            for k in range(len(hyperlink_text)):
                if len(hyperlink_text) > 0:
                    temp_text = doc.paragraphs[i].text.replace(hyperlink_text[k],
                                                               "[" + hyperlink_text[k] + "](" + hyperlink_url[k] + ")")
                    doc.paragraphs[i].text = temp_text
            nodes_temp = nodes_temp + doc.paragraphs[i].text.strip()
    nodes_text.append(nodes_temp)
    nodes_temp = ""

    if "Informations sur le cours" in nodes_text[0]:
        nodes_text[0] = nodes_text[0] + "Le code et le numéro du cours sont " + doc.paragraphs[0].text.strip() + ".\n"
        nodes_text[0] = nodes_text[0] + "Le titre du cours est " + doc.paragraphs[1].text.strip() + "."

    return nodes_text


def read_tables(html_text: str) -> List[pd.DataFrame]:
    soup = BeautifulSoup(html_text, 'html.parser')
    tables = soup.find_all('table')
    data_frames = []
    for table in tables:
        rows = table.find_all('tr')
        table_data = []
        for row in rows:
            cols = row.find_all('td')
            cols_text = [col.get_text() for col in cols]
            cols_links = [col.find('a')['href'] if col.find('a') and 'href' in col.find('a').attrs else '' for col in cols]
            cols_with_links = [f'[{text}] ({link})' if link else text for text, link in zip(cols_text, cols_links)]
            table_data.append(cols_with_links)
        df = pd.DataFrame(table_data)
        data_frames.append(df)
    return data_frames


def render_tables_add_to_nodes_text(nodes_text: List[str], doc_tables_df: List[pd.DataFrame]) -> List[str]:
    temp_text = ""
    for temp_df in doc_tables_df:
        temp_text = "*Informations*\n "
        for j in range(1, len(temp_df)):
            temp_text += f"Si vous êtes dans le tutoriel {temp_df.iloc[j, 0]}, votre instructeur responsable est {temp_df.iloc[j, 1]}.\n"
            temp_text += f"Si vous êtes dans le tutoriel {temp_df.iloc[j, 0]}, l'heure du tutoriel est {temp_df.iloc[j, 2]}.\n"
            temp_text += f"Si vous êtes dans le tutoriel {temp_df.iloc[j, 0]}, la salle du tutoriel est {temp_df.iloc[j, 3]}.\n"
            temp_text += f"Si vous êtes dans le tutoriel {temp_df.iloc[j, 0]}, l'adresse Zoom pendant les sessions en ligne est {temp_df.iloc[j, 4]}.\n"
        nodes_text.append(temp_text)

    return nodes_text


def clean_up(nodes_text: List[str]) -> List[str]:
    filtered_nodes_text = [text for text in nodes_text if not (text.endswith("*\n\n") or text.endswith("*\n \n"))]
    return filtered_nodes_text


def convert_to_json(sorted_nodes_text: List[str]) -> List[dict]:
    json_nodes_text = []
    for i, text in enumerate(sorted_nodes_text):
        node = {
            "node_number": i,
            "type": "NarrativeText",
            "text": text,
            "metadata": {
                "category_depth": 0,
                "page_number": 1,
                "languages": ["fr"],
                "filetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        }
        json_nodes_text.append(node)
    return json_nodes_text


def run_converter(docx_file_bytes: io.BytesIO) -> List[Element]:
    result = mammoth.convert_to_html(docx_file_bytes)

    doc = Document(docx_file_bytes)
    html_text = result.value

    sections = find_hlevel(doc)
    section_paragraphs = find_sections_paragraphs(sections, doc)
    nodes_text = convert_doc_to_nodes(section_paragraphs, doc, sections)
    doc_tables_df = read_tables(html_text)
    nodes_text = render_tables_add_to_nodes_text(nodes_text, doc_tables_df)
    sorted_nodes_text = clean_up(nodes_text)
    nodes = convert_to_json(sorted_nodes_text)
    elements = [Element(type=ElementType.of(node['type']), text=node['text'], metadata=node['metadata']) for node in nodes]

    return elements
