import io
import re  # regular expression
from typing import List

# Import packages ------------------------------------------------------------------------------------------ #
from docx import Document  # to import the syllabus (ocx file)
from docx.text.paragraph import Paragraph

from criaparse.models import ElementType, Element

section_length = 325 # Rough number of words in each section
# Functions -------------------------------------------------------------------------------------------------- #
def split_document(doc: Document) -> List[str]:
    paragraph_count = len(doc.paragraphs) # Number of paragraphs in document
    # Handle empty documents gracefully
    if paragraph_count == 0:
        return []

    section = []
    section_temp = doc.paragraphs[0].text # First paragraph in document
    for i in range(1, paragraph_count):
        next_temp = section_temp + " " + doc.paragraphs[i].text # next_temp includes the following paragraph
        # treat equality as acceptable to merge paragraphs whose combined
        # word count exactly matches the target section length
        if len(next_temp.split()) <= section_length: # the split is to have the number of words in next_temp
            section_temp = next_temp
        elif len(next_temp.split()) > section_length and (len(next_temp.split()) - section_length) > (section_length - len(section_temp.split())):
            # section with the augmented paragrah is bigger than section_length and the dela with section_length is bigger than the delta of the non-augmented section
            section.append(section_temp)
            section_temp =  doc.paragraphs[i].text # if next_temp is not included, the new section_temp is the current paragraph in the loop
        elif len(next_temp.split()) > section_length and (len(next_temp.split()) - section_length) < (section_length - len(section_temp.split())):
            section.append(next_temp)
            section_temp = ""

    section.append(section_temp) #append the last section_temp to the section list
    section_temp = ""

    return section

# Convert data to json
def convert_to_json(sections: List[str]) -> List[dict]:
    # First remove path from file_source
    node_number = 0
    json_nodes_text = []
    for text in sections:
        node = {
            "node_number": node_number,
            "type": "NarrativeText",
            "text": text,
            "metadata": {
                "category_depth": 0,
                "page_number": 1,  # not sure what to do with the page number here
                "languages": ["eng"],
                "filetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        }
        json_nodes_text.append(node)  # json.dumps converts the list to a JSON string
        node_number = node_number + 1
    return json_nodes_text

def run_converter(docx: io.BytesIO) -> List[Element]:

    doc = Document(docx)

    sections = split_document(doc)

    nodes: list[dict] = convert_to_json(sections)
    elements: list[Element] = []

    for node in nodes:
        elements.append(
            Element(
                type=ElementType.of(node['type']),
                text=node['text'],
                metadata=node['metadata']
            )
        )

    return elements
