# ---------------------------------------------------------------------------------------------------------- #
# Pre-processes a syllabus for Al the bot
# ---------------------------------------------------------------------------------------------------------- #
import io
import re  # regular expression
from typing import List

import mammoth  # to convert docx to html
import pandas as pd
from bs4 import BeautifulSoup
# Import packages ------------------------------------------------------------------------------------------ #
from docx import Document  # to import the syllabus (ocx file)

from criaparse.parser import Element, ElementType


# Functions -------------------------------------------------------------------------------------------------- #
def run_converter(docx: io.BytesIO) -> List[Element]:
    result = mammoth.convert_to_html(docx)

    doc = Document(docx)
    html_text = result.value

    def find_hlevel(doc):
        headings = []
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                headings.append(paragraph.text.strip())
        # Erase empty headers
        sections = [item for item in headings if
                    item]  # iterating item through headings, if item is true (i.e. not empty) add item in the list
        return sections

    def find_sections_paragraphs(sections, doc):  # finds the paragraph of the sections defined in the list sections
        section_paragraphs = []
        for i in range(len(sections)):
            for j in range(len(doc.paragraphs)):
                if sections[i] == doc.paragraphs[
                    j].text.strip():  # if the section in the list is the same as the doc paragraph, add the paragraph number to section_paragraph list
                    section_paragraphs.append(j)
        return section_paragraphs

    def convert_doc_to_nodes(section_paragraphs, doc):
        nodes_text = []
        nodes_temp = ""
        for i in range(len(section_paragraphs) - 1):
            for j in range(section_paragraphs[i] + 1, section_paragraphs[i + 1]):
                # first check if there's a hyperlink in the paragraph
                hyperlink_text, hyperlink_url = include_hyperlink(doc.paragraphs[j])
                for k in range(len(hyperlink_text)):  # loop to grab all items in hyperlink_text[] and hyperlink_url[]
                    if len(hyperlink_text) > 0:  # if there's no item in hypertext_link or hypertext_url, you get an error message
                        temp_text = doc.paragraphs[j].text.replace(hyperlink_text[k],
                                                                   "[" + hyperlink_text[k] + "](" + hyperlink_url[
                                                                       k] + ")")  # follows the mark down format
                        doc.paragraphs[j].text = temp_text
                nodes_temp = nodes_temp + doc.paragraphs[j].text.strip() + " "
            nodes_text.append("*" + sections[i] + "*\n" + nodes_temp + "\n")
            nodes_temp = ""
        # need to add the following code to capture the paragraphs in the last section, which are not captured in the loop because the loop would be out of range
        for i in range(section_paragraphs[len(section_paragraphs) - 1],
                       len(doc.paragraphs)):  # from the last section position to the last paragraph of the document
            if i == section_paragraphs[len(section_paragraphs) - 1]:
                nodes_temp = nodes_temp + "*" + doc.paragraphs[i].text.strip() + "*\n"
            else:
                # first check if there's a hyperlink in the paragraph
                hyperlink_text, hyperlink_url = include_hyperlink(doc.paragraphs[i])
                for k in range(len(hyperlink_text)):  # loop to grab all items in hyperlink_text[] and hyperlink_url[]
                    if len(hyperlink_text) > 0:  # if there's no item in hypertext_link or hypertext_url, you get an error message
                        temp_text = doc.paragraphs[i].text.replace(hyperlink_text[k],
                                                                   "[" + hyperlink_text[k] + "](" + hyperlink_url[
                                                                       k] + ")")  # follows the mark down format
                        doc.paragraphs[i].text = temp_text
                nodes_temp = nodes_temp + doc.paragraphs[i].text.strip()
        nodes_text.append(nodes_temp)
        nodes_temp = ""

        # Add the course title, rubric and number, which are not in Course Information, but in the title
        if "Course Information" in nodes_text[0]:  # The following code does not apply with Questions.docx
            nodes_text[0] = nodes_text[0] + "The course rubric and number is " + doc.paragraphs[0].text.strip() + ".\n"
            nodes_text[0] = nodes_text[0] + "The course title is " + doc.paragraphs[1].text.strip() + "."

        return nodes_text

    def read_tables_bs4mp(html_text):
        # Parse the HTML with BeautifulSoup
        soup = BeautifulSoup(html_text, 'html.parser')

        # Find all tables
        tables = soup.find_all('table')

        data_frames = []
        for table in tables:
            # Find all rows
            rows = table.find_all('tr')

            table_data = []
            for row in rows:
                cols = row.find_all('td')

                # Extract text and links
                cols_text = [col.get_text() for col in cols]
                cols_links = [col.find('a')['href'] if col.find('a') and 'href' in col.find('a').attrs else '' for col
                              in
                              cols]

                # Combine text and links
                cols_with_links = [f'[{text}] ({link})' if link else text for text, link in zip(cols_text, cols_links)]

                table_data.append(cols_with_links)

            # Convert table data to DataFrame
            df = pd.DataFrame(table_data)

            data_frames.append(df)

        return data_frames

    def read_tables(html_text):
        # With the following two code lines, pd.read_html did not keep the links. So I used beautiful soup instead
        # html_buffer = StringIO(html_text)  # Had to wrap the HTML string in a StringIO object because a direct pd.read (see line below) will be deprecated
        # doc_tables_df = pd.read_html(html_buffer)  # This grabs all the tables in the syllabus and stores them in a dataframe
        doc_tables_df = read_tables_bs4mp(html_text)

        # find all caption
        # pattern = r"</table>(.*?)</p>"  # All captions are right after the table between <p> and </p>
        # matches = re.findall(pattern, html_text)
        # table_titles = []
        # for i in range(len(matches)):
        #     if matches[i][5].isdigit():
        #         table_titles.append(matches[i][5:])  # Create table title without the initial 2 digits
        #     else:
        #         table_titles.append(matches[i][4:])

        # find name of table, which is the heading above it (h1, h2 or h3
        longest_section_len = len(max(sections, key=len))  # first determine the length of the longest section
        pattern_close = r"</h\d>\s*<table>"  # Finds all h1 or h2 or h3 followed by a table
        closing_indexes = [m.start() for m in
                           re.finditer(pattern_close, html_text)]  # Finds the closing index of the table

        previous_index = []  # this goes back to the index at a distance of the longest section
        for closing_index in closing_indexes:
            previous_index.append(closing_index - longest_section_len)
        chunks = []  # chunks contain the titles but addtional crap beforehand
        for i in range(len(closing_indexes)):
            chunks.append(html_text[previous_index[i]:closing_indexes[i] + 12])

        almost_titles = []  # an almost title still needs to be cleaned further to get the title
        for item in chunks:
            match = re.search(r"<h\d>.*</h", item)
            if match:
                almost_titles.append(match.group())
        # Clean the almost titles to have real titles
        table_titles = []
        for i in range(len(almost_titles)):  # Removes the initial tag and the ending tag
            table_titles.append(almost_titles[i][4:-3])

        for i in range(
                len(table_titles)):  # this part is to find occurrence of another heading in the title (one that should be eliminated)
            if "</h" in table_titles[i]:
                temp_index = table_titles[i].find("</h") + 9
                table_titles[i] = table_titles[i][temp_index:]

        return doc_tables_df, table_titles

    def render_tables_add_to_nodes_text(table_titles, nodes_text, doc_tables_df):
        i = 0
        temp_text = ""
        for title in table_titles:
            temp_df = (doc_tables_df[i])

            if title == "Tutorials":
                temp_text = "*Tutorials*\n "
                for j in range(1, len(temp_df)):
                    temp_text = temp_text + "If you are in Tutorial " + temp_df.iloc[
                        j, 0] + ", your TA (or responsible instructor who teaches the tutorial) is " + temp_df.iloc[
                                    j, 1] + ".\n "
                    temp_text = temp_text + "If you are in Tutorial " + temp_df.iloc[
                        j, 0] + ", your tutorial time is " + \
                                temp_df.iloc[j, 2] + ".\n "
                    temp_text = temp_text + "If you are in Tutorial " + temp_df.iloc[
                        j, 0] + ", your tutorial room is " + \
                                temp_df.iloc[j, 3] + ".\n "
                    temp_text = temp_text + "If you are in Tutorial " + temp_df.iloc[
                        j, 0] + ", your Zoom address during online sessions is " + temp_df.iloc[j, 4] + ".\n"
                nodes_text.append(temp_text)

            elif title == "Faculty Members Information":
                temp_text = "*Faculty Members Information*\n "
                for j in range(1, len(temp_df)):
                    temp_text = temp_text + temp_df.iloc[j, 0] + " is the course's " + temp_df.iloc[
                        j, 1] + " and has the following email address: " + temp_df.iloc[
                                    j, 2] + " and has the following office hours (time you can meet or appointment time): " + \
                                temp_df.iloc[
                                    j, 3] + "and has the following office address or location (where you can meet with your professor or instructor or teacher or TA): " + \
                                temp_df.iloc[j, 4] + ".\n "
                nodes_text.append(temp_text)

            elif title == "Summary of Evaluation":
                temp_text = "*Summary of Evaluation*\n "
                temp_text = temp_text + "This section answers questions about how much an assignment is worth (how much it counts toward the final grade) and when the assignments are due or have to be submitted or handed in (submission date). \n"
                # evaluation_table_index = i  # this is to be able to find the evaluation table when we do the query precoessing for temporal relations
                for j in range(1, len(temp_df)):
                    temp_text = temp_text + "The " + temp_df.iloc[j, 0] + " is worth " + temp_df.iloc[
                        j, 1] + " of the final grade. In other words, it counts for " + temp_df.iloc[
                                    j, 1] + " of the final grade.\n "
                    temp_text = temp_text + "The " + temp_df.iloc[j, 0] + " is due on " + temp_df.iloc[
                        j, 2] + ". In other words, the deadline or due date or submission date for " + temp_df.iloc[
                                    j, 0] + " is " + temp_df.iloc[j, 2] + ".\n "
                nodes_text.append(temp_text)

            elif title == "Grading Equivalence":
                temp_text = "*Grading Equivalence*\n "
                for j in range(1, len(temp_df)):
                    temp_text = temp_text + temp_df.iloc[j, 0] + " is the same as a grade point of " + temp_df.iloc[
                        j, 1] + ", which falls in the percent range of " + temp_df.iloc[
                                    j, 2] + "%, and is described as '" + \
                                temp_df.iloc[j, 3] + "'.\n "
                nodes_text.append(temp_text)

            elif title == "Definitions of Standing":
                temp_text = "*Definitions of Standing*\n "
                for j in range(0, len(temp_df)):
                    temp_text = temp_text + "A grade considered '" + temp_df.iloc[j, 0] + "' means that you have a " + \
                                temp_df.iloc[j, 1] + "\n "
                nodes_text.append(temp_text)

            elif title == "Schedule and Readings":
                temp_text = "*Schedule and Readings*\n "
                for j in range(1, len(temp_df)):
                    temp_text = temp_text + "The topic on " + temp_df.iloc[j, 2] + " is '" + temp_df.iloc[
                        j, 0] + "'. In other words, '" + temp_df.iloc[j, 0] + "' is presented on " + temp_df.iloc[
                                    j, 2] + ".\n "
                    if str(temp_df.iloc[j, 1]) == "nan":
                        temp_text = temp_text + "There are no readings on " + temp_df.iloc[j, 2] + ".\n "
                    else:
                        temp_text = temp_text + "The reading(s) for the topic called '" + temp_df.iloc[j, 0] + "' on " + \
                                    temp_df.iloc[j, 2] + " is (are) the following: " + str(temp_df.iloc[j, 1]) + "\n "
                nodes_text.append(temp_text)

            elif title == "Important Dates":
                temp_text = "*Important Dates*\n "
                for j in range(1, len(temp_df)):
                    if "None" in temp_df.iloc[j, 1]:
                        temp_text = temp_text + "There is no " + temp_df.iloc[j, 0] + ".\n "
                    else:
                        temp_text = temp_text + temp_df.iloc[j, 0] + " is on " + temp_df.iloc[j, 1] + ".\n "
                nodes_text.append(temp_text)

            else:
                temp_text = "*" + title + "*\n "
                nb_rows = len(temp_df)
                nb_columns = len(temp_df.columns)
                for j in range(1, nb_rows):
                    temp_text = temp_text + "The following " + temp_df.iloc[0, 0].lower() + ": " + temp_df.iloc[
                        j, 0] + " has "
                    for k in range(1, nb_columns - 1):
                        temp_text = temp_text + "the following " + temp_df.iloc[0, k].lower() + ": " + str(
                            temp_df.iloc[j, k]) + " and has "
                    temp_text = temp_text + "the following " + temp_df.iloc[0, k + 1].lower() + ": " + str(
                        temp_df.iloc[j, k + 1]).strip() + "."
                nodes_text.append(temp_text)
            i = i + 1

        return nodes_text

    def clean_up(nodes_text, sections):
        # Render the section Course Information
        temp_text = nodes_text[0].replace("Course Director:",
                                          "The course director (or professor or instructor or teacher) for this course is ")
        nodes_text[0] = temp_text
        temp_text = nodes_text[0].replace("Email:", "\n Your course director's email is ")
        nodes_text[0] = temp_text
        temp_text = nodes_text[0].replace("Semester:", "\n The current semester (or term) is ")
        nodes_text[0] = temp_text
        temp_text = nodes_text[0].replace("Lecture time & day:",
                                          "\n The lecture (or class) is offered on the following day and time: ")
        nodes_text[0] = temp_text
        temp_text = nodes_text[0].replace("Lecture room:",
                                          "\n If you're wondering how to get to your lecture classroom, the lecture (or class) takes place in the following classroom (or location): ")
        nodes_text[0] = temp_text
        temp_text = nodes_text[0].replace("Zoom (Lecture):",
                                          "\n Some classes may be offered on Zoom or you may have to attend some classes on Zoom only during unforeseen situations such as snowstorms or the instructor's illness, in which case the Zoom link (or Zoom address) for the lecture will be ")
        nodes_text[0] = temp_text
        temp_text = nodes_text[0].replace("eClass:",
                                          "\n There is an eClass site (the course has been uploaded to eClass) and the eClass link (or address or URL) is ")
        nodes_text[0] = temp_text
        temp_text = nodes_text[0].replace("Office:",
                                          "\n What is the course director's (or professor's or instructor's or teacher's) office number (or office address)? Where can I meet him or her? The answer is: ")
        nodes_text[0] = temp_text
        temp_text = nodes_text[0].replace("Office Hours:",
                                          "\n The course director's (or professor's or instructor's or teacher's) office hours are ")
        nodes_text[0] = temp_text
        temp_text = nodes_text[0].replace("\t", "")
        nodes_text[0] = temp_text

        # Combine "Tutorials" and "Faculty Members Information" for better results
        tutorials_index = 0
        fac_memberss_index = 0
        for i in range(len(nodes_text)):
            if nodes_text[i].find("*Tutorials*") > -1:
                tutorials_index = i
            if nodes_text[i].find("*Faculty Members Information*") > -1:
                fac_memberss_index = i
        if tutorials_index != 0 and fac_memberss_index != 0:
            nodes_text[tutorials_index] = nodes_text[tutorials_index] + nodes_text[fac_memberss_index]
            del nodes_text[fac_memberss_index]

        # Erase empty nodes
        filtered_nodes_text = [text for text in nodes_text if not (text.endswith("*\n\n") or text.endswith("*\n \n"))]
        nodes_text = filtered_nodes_text

        # Combine two nodes when there is a table and text under the same header
        sorted_nodes_text = sorted(nodes_text)  # If you sort the list, like item will be next to each other
        for i in range(len(nodes_text) - 2, -1, -1):
            first_index = sorted_nodes_text[i].find("*")
            second_index = sorted_nodes_text[i].find("*", first_index + 1)  # Starts searching after the first *
            temp_node = sorted_nodes_text[i][:second_index]  # you now have the title of the node
            if temp_node == sorted_nodes_text[i + 1][:second_index]:
                sorted_nodes_text[i] = sorted_nodes_text[i] + sorted_nodes_text[i + 1][
                                                              second_index + 1:]  # Combine the following item with the previous
                del sorted_nodes_text[i + 1]  # And delete the following, now redundant

        # From here on (i.e. after clean_up), we must work with sort_nodes_text instead of nodes_text

        # Erase the caption that appears after the title in sorted_nodes_text
        # i = 0
        # for node in sorted_nodes_text:
        #     first_index = node.find("*")
        #     second_index = node.find("*", first_index + 1)  # Starts searching after the first *
        #     title_length = second_index - first_index - 1
        #     initial_title = node[first_index+1:second_index]
        #     potential_caption_index = node.find(initial_title, second_index, second_index + title_length + 4)
        #     potential_caption = node[potential_caption_index:potential_caption_index + title_length]
        #     if initial_title == potential_caption:
        #         node = "*" + initial_title + "*\n" + node[potential_caption_index + title_length + 1:]
        #     sorted_nodes_text[i] = node
        #     i = i + 1
        return sorted_nodes_text

    def include_hyperlink(
            paragraph):  # This function looks for hyperlinks in a paragraph. If found, returns the list of text that has a link and the list of its url
        # has_hyperlink = False
        hyperlink_text = []
        hyperlink_url = []
        if len(paragraph.hyperlinks) > 0:

            for hyperlink in paragraph.hyperlinks:
                hyperlink_text.append(hyperlink.text)
                hyperlink_url.append(hyperlink.url)

        return hyperlink_text, hyperlink_url

    def convert_to_json(sorted_nodes_text) -> list[dict]:
        # First remove path from file_source
        node_number = 0
        json_nodes_text = []
        for text in sorted_nodes_text:
            node = {
                "node_number": node_number,
                "type": "NarrativeText",
                "text": text,
                "metadata": {
                    "category_depth": 0,
                    "page_number": 1,  # not sure what to do with the page number here
                    "languages": ["eng"],
                    "filetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                }
            }
            json_nodes_text.append(node)  # json.dumps converts the list to a JSON string
            node_number = node_number + 1
        return json_nodes_text

    sections = find_hlevel(doc)
    section_paragraphs = find_sections_paragraphs(sections,
                                                  doc)  # the sections in the sections list are assigned a paragraph
    nodes_text = convert_doc_to_nodes(section_paragraphs,
                                      doc)  # the doc is converted to a list of semantic sections containing the text
    doc_tables_df, table_titles = read_tables(
        html_text)  # I need the dataframe created in read_table to use in render_tables_add_to_notes, where the dataframe is rendered
    # read_tables_temp(html_text)
    render_tables_add_to_nodes_text(table_titles, nodes_text,
                                    doc_tables_df)  # Where the rendering of tables is done and added to the list nodes_text
    sorted_nodes_text = clean_up(nodes_text, sections)  # final touches to clean up the list

    nodes: list[dict] = convert_to_json(sorted_nodes_text)
    elements: list[Element] = []

    for node in nodes:
        elements.append(
            Element(
                type=ElementType.of(node['type']),
                text=node['text'],
                metadata=node['metadata']
            )
        )

    return elements
